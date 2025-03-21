import requests
from bs4 import BeautifulSoup
import json
import re
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import TimeoutException, WebDriverException
import os
import time
import random
import docx
from pdfminer.high_level import extract_text
import subprocess
from webdriver_manager.chrome import ChromeDriverManager
import logging
import concurrent.futures
from functools import partial
import sys

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.FileHandler("scraper.log"), logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

def extract_text_from_document(filepath, output_dir):
    """
    Extract text from a document based on its file type.
    This function is separated to be called on-demand.
    """
    try:
        extracted_text = ""
        if filepath.lower().endswith('.docx'):
            extracted_text = extract_text_from_docx(filepath)
        elif filepath.lower().endswith('.pdf'):
            extracted_text = extract_text_from_pdf(filepath)
        elif filepath.lower().endswith('.doc'):
            # Try LibreOffice first (more reliable than antiword)
            try:
                temp_output = os.path.join(output_dir, "temp_converted.txt")
                command = ["soffice", "--headless", "--convert-to", "txt:Text", "--outdir", output_dir, filepath]
                result = subprocess.run(command, capture_output=True, text=True, timeout=60)
                if os.path.exists(temp_output):
                    with open(temp_output, 'r', encoding='utf-8', errors='replace') as f:
                        extracted_text = f.read()
                    os.remove(temp_output)
            except (subprocess.SubprocessError, FileNotFoundError):
                # Fallback to antiword
                command = ["antiword", filepath]
                result = subprocess.run(command, capture_output=True, text=True, timeout=30)
                extracted_text = result.stdout
        elif filepath.lower().endswith('.txt'):
            with open(filepath, 'r', encoding='utf-8', errors='replace') as txt_file:
                extracted_text = txt_file.read()
        
        return extracted_text
    except FileNotFoundError:
        return "Error: No conversion tool available. Please install LibreOffice or antiword."
    except subprocess.SubprocessError as e:
        return f"Error extracting text: {e}"
    except Exception as e:
        logger.error(f"Unexpected error extracting text from {filepath}: {e}")
        return ""

def download_and_process_doc(doc_info, output_dir, extract_text=False):
    """
    Downloads a document and optionally extracts text.
    
    Parameters:
    - doc_info: Dictionary with document information
    - output_dir: Directory to save files
    - extract_text: If True, extract text from the document; if False, just download
    """
    doc_url = doc_info["link_url"]
    filename = os.path.basename(doc_url)
    
    # Skip User-Manual files
    if "User-Manual" in doc_url or "User-Manual" in filename:
        logger.info(f"Skipping User-Manual file: {doc_url}")
        return doc_info
    
    filepath = os.path.join(output_dir, filename)
    text_filepath = os.path.join(output_dir, filename + ".txt")
    doc_info["filepath"] = filepath
    doc_info["text_filepath"] = text_filepath
    doc_info["downloaded"] = os.path.exists(filepath)
    doc_info["text_extracted"] = os.path.exists(text_filepath)

    # If both files exist and we need text extraction, load the cached text
    if os.path.exists(filepath) and os.path.exists(text_filepath) and extract_text:
        logger.info(f"Using cached text for: {filename}")
        try:
            with open(text_filepath, 'r', encoding='utf-8', errors='replace') as f:
                doc_info['extracted_text'] = f.read()
        except Exception as e:
            logger.warning(f"Error reading cached text for {filename}: {e}")
            doc_info['extracted_text'] = ""
        return doc_info

    # If file exists but we don't need text, just return the info
    if os.path.exists(filepath) and not extract_text:
        logger.info(f"File exists, skipping download: {filename}")
        return doc_info

    # Use backoff strategy for downloads
    max_retries = 3
    for retry in range(max_retries):
        try:
            logger.info(f"Downloading: {doc_url}")
            # Use a session with appropriate headers
            session = requests.Session()
            session.headers.update({
                'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.5',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1',
                'Sec-Fetch-Dest': 'document',
                'Sec-Fetch-Mode': 'navigate',
                'Sec-Fetch-Site': 'none',
                'Sec-Fetch-User': '?1',
                'Pragma': 'no-cache',
                'Cache-Control': 'no-cache',
            })
            
            response = session.get(doc_url, stream=True, timeout=60)
            response.raise_for_status()

            with open(filepath, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            
            doc_info["downloaded"] = True

            # Only extract text if requested
            if extract_text:
                extracted_text = extract_text_from_document(filepath, output_dir)
                
                if extracted_text:
                    with open(text_filepath, 'w', encoding='utf-8') as tf:
                        tf.write(extracted_text)
                    logger.info(f"Extracted text saved to: {text_filepath}")
                    doc_info['extracted_text'] = extracted_text
                    doc_info["text_extracted"] = True
                else:
                    logger.warning(f"No text extracted from {filename}")
                    doc_info['extracted_text'] = ""
                    doc_info["text_extracted"] = False

            # Successful download, break the retry loop
            break
            
        except requests.exceptions.RequestException as e:
            if retry < max_retries - 1:
                wait_time = (2 ** retry) * 5  # Exponential backoff
                logger.warning(f"Error downloading {doc_url}: {e}. Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
            else:
                logger.error(f"Failed to download {doc_url} after {max_retries} attempts: {e}")
                doc_info['error'] = f"Error downloading {doc_url}: {e}"
                doc_info["downloaded"] = False
        except (OSError, IOError) as e:
            logger.error(f"Error saving or processing {doc_url}: {e}")
            doc_info['error'] = f"Error saving or processing {doc_url}: {e}"
        except Exception as e:
            logger.error(f"An unexpected error occurred processing {doc_url}: {e}")
            doc_info['error'] = f"An unexpected error occurred processing {doc_url}: {e}"
    
    return doc_info

def scrape_and_download(url, output_dir="scraped_data"):
    """Scrapes structured data and downloads/extracts text from documents."""

    from selenium.webdriver.chrome.service import Service as ChromeService

    # --- 1. Selenium Setup (Enhanced Robustness) ---
    options = Options()
    options.add_argument("--headless=new")  # Updated headless mode
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--window-size=1920,1080")
    options.add_argument("--disable-gpu")
    options.add_argument("--disable-extensions")
    # Add user agent to appear more like a normal browser
    options.add_argument("--user-agent=Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
    # Configure Chrome to block images and other resources
    options.add_argument('--blink-settings=imagesEnabled=false')
    
    # Add proxy if needed (uncomment if you have a proxy service)
    # options.add_argument('--proxy-server=http://your-proxy-server:port')
    
    driver = None
    try:
        logger.info("Installing ChromeDriver...")
        driver_path = ChromeDriverManager().install()
        
        logger.info("Creating ChromeService...")
        service = ChromeService(executable_path=driver_path)
        
        logger.info("Initializing Chrome WebDriver...")
        driver = webdriver.Chrome(service=service, options=options)
        
        logger.info("Setting page load timeout...")
        driver.set_page_load_timeout(180)  # 3 minutes timeout
        
        logger.info(f"Loading URL: {url}")
        driver.get(url)
        
        # Wait for page to load with exponential backoff
        max_attempts = 5
        base_wait_time = 10
        
        # Inside your function
        for attempt in range(max_attempts):
            try:
                logger.info(f"Waiting for elements to load (attempt {attempt+1}/{max_attempts})...")
                
                WebDriverWait(driver, base_wait_time * (2 ** attempt)).until(
                    lambda driver: (
                        len(driver.find_elements(By.CLASS_NAME, "text-2xl")) > 0 or
                        len(driver.find_elements(By.CLASS_NAME, "mt-12")) > 0 or
                        len(driver.find_elements(By.XPATH, "//li[contains(@class, 'relative flex justify-between')]")) > 0
                    )
                )
                
                logger.info("Page loaded successfully.")
                break
            except TimeoutException:
                if attempt < max_attempts - 1:
                    logger.warning(f"Timeout waiting for page to load. Retrying...")
                    # Take a screenshot to see what's happening
                    driver.save_screenshot(f"loading_attempt_{attempt+1}.png")
                    # Refresh the page and try again
                    driver.refresh()
                else:
                    raise
        
        # Add a short delay even after elements are found to ensure JavaScript has run
        time.sleep(5)
        
        page_source = driver.page_source
        
        # Save the page source for debugging
        with open("page_source.html", "w", encoding="utf-8") as f:
            f.write(page_source)
        
        logger.info("Page source saved to page_source.html")

    except Exception as e:
        try:
            if driver:
                driver.save_screenshot("error_screenshot.png")
                logger.error(f"Selenium error: {e}. Screenshot saved to error_screenshot.png")
        except:
            logger.error(f"Selenium error: {e}. Could not save screenshot.")
        
        try:
            if driver:
                driver.quit()
        except:
            pass
            
        return {"error": f"Selenium error: {str(e)}"}

    # --- 2. BeautifulSoup Parsing (Structured Data) ---
    soup = BeautifulSoup(page_source, 'html.parser')

    # Check if we got a meaningful page
    if len(page_source) < 1000 or "Access Denied" in page_source:
        logger.error("Page access denied or returned minimal content")
        if driver:
            driver.quit()
        return {"error": "Page access denied or returned minimal content"}

    data = {
        "measure_number": None,
        "title": None,
        "status": None,
        "filing_date": None,
        "authors": [],
        "origin_chamber": None,
        "current_chamber": None,
        "topic": None,
        "other_data": {},
        "documents": [],
        "tramites": [],
        "votaciones": [],
        "comisiones": []
    }

    # Extract measure number from heading
    header = soup.find('h1', class_=lambda c: c and "text-2xl" in c)
    if header:
        data["measure_number"] = header.get_text(strip=True)
        logger.info(f"Found measure number: {data['measure_number']}")
    
    # Extract filing date
    filing_date_elem = soup.find('span', string=lambda s: s and "Fecha de Radicación" in s)
    if filing_date_elem:
        date_span = filing_date_elem.find_next('span', class_=lambda c: c and "text-xs" in c)
        if date_span:
            data["filing_date"] = date_span.get_text(strip=True)
            logger.info(f"Found filing date: {data['filing_date']}")
    
    # Extract title
    title_elem = soup.find('span', string=lambda s: s and "Título" in s)
    if title_elem:
        # Title is usually in a sibling or nearby span with text-balance class
        title_span = title_elem.find_next('span', class_="text-balance")
        if title_span:
            data["title"] = title_span.get_text(strip=True)
            logger.info(f"Found title: {data['title'][:50]}...")
    
    # Extract authors from the "Autores" tab, if possible
    authors_section = soup.find(lambda tag: tag.name == 'button' and tag.get_text(strip=True) == 'Autores (1)')
    if authors_section:
        author_items = soup.find_all('li', class_=lambda c: c and "autor_id_li" in c)
        for author_item in author_items:
            author_name = author_item.find('p', class_=lambda c: c and "font-semibold" in c)
            if author_name:
                data["authors"].append(author_name.get_text(strip=True))
                logger.info(f"Found author: {author_name.get_text(strip=True)}")
    
    # Fallback to direct author parsing if needed
    if not data["authors"]:
        authors_div = soup.find(lambda tag: tag.name == 'div' and tag.get_text(strip=True) == 'Autores')
        if authors_div:
            author_spans = authors_div.find_next('div').find_all('span')
            for span in author_spans:
                if span.get_text(strip=True) and not any(kw in span.get_text().lower() for kw in ["autor", "fecha"]):
                    data["authors"].append(span.get_text(strip=True))
                    logger.info(f"Found author: {span.get_text(strip=True)}")

    # --- Helper Function (Extract Documents) ---
    def extract_documents_from_element(element):
        documents = []
        if element:
            for link in element.find_all('a', href=True):
                doc_url = link['href']
                if not doc_url.startswith("http"):
                    doc_url = "https://sutra.oslpr.org" + doc_url
                doc_description = link.get_text(strip=True)
                documents.append({"link_url": doc_url, "description": doc_description})
        return documents

    # --- 3. Extract Commission Information ---
    # In the modern UI, commission info is integrated into the events
    # However, we can check if "Comisión" tab exists
    commission_tab = soup.find(lambda tag: tag.name == 'button' and 'Comisión' in tag.get_text())
    if commission_tab:
        logger.info("Found commission tab")
        # Extract commission info from the tab content
        # This will depend on the exact structure, but we might find it in a list
        commission_items = soup.find_all('li', string=lambda s: s and "Comisión" in s)
        for item in commission_items:
            commission_name = item.get_text(strip=True)
            if commission_name:
                data["comisiones"].append({
                    "comision": commission_name,
                    "documents": []  # We'll have to find the documents elsewhere
                })
    
    # --- 4. Extract Documents Tab Content ---
    # Find the "Documentos" tab and its content
    documents_tab = soup.find(lambda tag: tag.name == 'button' and tag.get_text(strip=True) == 'Documentos (0)')
    if documents_tab:
        logger.info("Found documents tab")
        # Documents might be in a following div or list
        document_items = soup.find_all('a', href=lambda h: h and (h.endswith('.pdf') or h.endswith('.doc') or h.endswith('.docx')))
        for doc_link in document_items:
            doc_url = doc_link['href']
            if not doc_url.startswith("http"):
                doc_url = "https://sutra.oslpr.org" + doc_url
            
            # Try to find document description
            doc_desc_elem = doc_link.find('span', class_=lambda c: c and "text-sutra-secondary" in c)
            doc_desc = doc_desc_elem.get_text(strip=True) if doc_desc_elem else doc_link.get_text(strip=True)
            
            # Only add if not already in any event
            if not any(doc['link_url'] == doc_url for tramite in data['tramites'] for doc in tramite['documents']):
                data["documents"].append({
                    "link_url": doc_url,
                    "description": doc_desc
                })
    
    # Additional document extraction from all links
    # This ensures we don't miss documents that might not be in the tab
    for doc_link in soup.find_all('a', href=lambda h: h and (h.endswith('.pdf') or h.endswith('.doc') or h.endswith('.docx'))):
        doc_url = doc_link['href']

        # Skip User-Manual files
        if "User-Manual" in doc_url:
            logger.info(f"Skipping User-Manual file in extraction: {doc_url}")
            continue

        if not doc_url.startswith("http"):
            doc_url = "https://sutra.oslpr.org" + doc_url
        
        # Try to find document description
        doc_span = doc_link.find('span', class_=lambda c: c and "cursor-pointer" in c)
        doc_desc = doc_span.get_text(strip=True) if doc_span else doc_link.get_text(strip=True)
        
        # Only add if not already in any list
        if not any(doc['link_url'] == doc_url for doc_list in [data['documents']] + [tramite['documents'] for tramite in data['tramites']] for doc in doc_list):
            data["documents"].append({
                "link_url": doc_url,
                "description": doc_desc
            })
    
    logger.info(f"Found {len(data['documents'])} general documents")

    # --- 5. Extract Events from the modern UI structure ---
    event_items = soup.find_all('li', class_=lambda c: c and "relative flex justify-between" in c)
    logger.info(f"Found {len(event_items)} event items")

    for event_item in event_items:
        # Extract event title/type
        event_title_elem = event_item.find('span', class_="text-sutra-primary")
        if not event_title_elem:
            continue
            
        event_title = event_title_elem.get_text(strip=True)
        
        # Initialize event data
        event_data = {
            "descripcion": event_title,
            "fecha": None,
            "documents": []
        }
        
        # Extract date
        date_elem = event_item.find('span', string=lambda s: s and "Fecha:" in s)
        if date_elem:
            date_parent = date_elem.parent
            # The date text is the text after "Fecha: "
            event_data["fecha"] = date_parent.get_text(strip=True).replace("Fecha:", "").strip()
        
        # Extract commission information - this needs to find the actual commission name
        # This looks for the last <p> element with text in the event item, which typically contains the commission name
        all_paragraphs = event_item.find_all('p', class_="mt-1 flex text-xs leading-5 text-gray-500")
        for p in all_paragraphs:
            # Skip paragraphs that contain "Fecha:" or document links
            if p.find('span', string=lambda s: s and "Fecha:" in s) or p.find('a'):
                continue
            
            # Check if this paragraph contains text about a commission
            commission_text = p.get_text(strip=True)
            if commission_text and "Comisión" in commission_text:
                # This is likely the commission name
                event_data["comision"] = commission_text
                break
        
        # Extract documents
        doc_links = event_item.find_all('a', href=True)
        for link in doc_links:
            doc_url = link['href']
            # Skip User-Manual files
            if "User-Manual" in doc_url:
                logger.info(f"Skipping User-Manual file in extraction: {doc_url}")
                continue
                
            # Make sure URL is absolute
            if not doc_url.startswith("http"):
                doc_url = "https://sutra.oslpr.org" + doc_url
                
            # Try to find document description
            doc_desc_elem = link.find('span', class_=lambda c: c and "text-sutra-secondary" in c)
            doc_desc = doc_desc_elem.get_text(strip=True) if doc_desc_elem else "Document"
            
            event_data["documents"].append({
                "link_url": doc_url,
                "description": doc_desc
            })
        
        # Determine if this is a vote event
        if "Votación" in event_title or "Aprobado" in event_title:
            # Try to extract vote counts
            vote_counts = {}
            vote_elements = event_item.find_all('span', string=lambda s: s and any(x in s for x in ["Votos a favor", "Votos en contra", "Votos abstenidos", "Votos ausentes"]))
            
            for vote_elem in vote_elements:
                vote_type = vote_elem.get_text(strip=True).rstrip(":")
                vote_value_elem = vote_elem.parent
                if vote_value_elem:
                    try:
                        # Extract just the number
                        vote_text = vote_value_elem.get_text(strip=True).replace(vote_type, "").strip()
                        vote_counts[vote_type] = int(vote_text) if vote_text.isdigit() else vote_text
                    except:
                        pass
            
            data["votaciones"].append({
                "camara": "Senado" if "Senado" in event_title else "Cámara",
                "fecha": event_data["fecha"],
                "descripcion": event_title,  # Changed from "resultado" to "descripcion"
                "documents": event_data["documents"],
                "votes": vote_counts if vote_counts else None
            })
        else:
            # This is a regular "tramite" (processing step)
            data["tramites"].append(event_data)
    
    logger.info(f"Extracted {len(data['tramites'])} tramites and {len(data['votaciones'])} votaciones")

    if driver:
        driver.quit()

    # --- 7. Prepare Documents (with optional downloading) ---
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    all_documents = data["documents"] + [doc for tramite in data["tramites"] for doc in tramite["documents"]] + [doc for votacion in data["votaciones"] for doc in votacion["documents"]] + [doc for comision in data["comisiones"] for doc in comision["documents"]]

    logger.info(f"Found {len(all_documents)} documents")

    # Check if we're applying the monkey patch 
    if download_and_process_doc.__name__ == 'patched_download_func':
        logger.info("IMPORTANT: Document downloads SKIPPED due to --no-extract flag")
        processed_documents = [download_and_process_doc(doc, output_dir) for doc in all_documents]
        # Add timing info
        logger.info(f"Processed {len(processed_documents)} document metadata in NO DOWNLOAD mode")
    else:
        # Use ThreadPoolExecutor to download files in parallel WITHOUT text extraction
        logger.info(f"Starting parallel download of {len(all_documents)} documents (no text extraction)")
        processed_documents = []
        with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
            # Pass extract_text=False to skip text extraction
            download_fn = partial(download_and_process_doc, output_dir=output_dir, extract_text=False)
            # Process results as they complete to track errors
            for i, doc_info in enumerate(executor.map(download_fn, all_documents)):
                processed_documents.append(doc_info)

                # If document had an error, add it to the data errors list
                if 'error' in doc_info:
                    data.setdefault("errors", []).append(doc_info['error'])
                    del doc_info['error']  # Remove temporary error field
                    
        logger.info(f"Completed downloading {len(processed_documents)} documents")

    # Replace the all_documents list with the processed one
    all_documents = processed_documents

    # Combine tramites and votaciones into a single eventos array
    eventos = []

    # Add all tramites to eventos
    for tramite in data["tramites"]:
        evento = tramite.copy()
        evento["tipo"] = "tramite"
        eventos.append(evento)

    # Add all votaciones to eventos
    for votacion in data["votaciones"]:
        evento = votacion.copy()
        evento["tipo"] = "votacion"
        # Rename resultado to descripcion
        if "resultado" in evento:
            evento["descripcion"] = evento.pop("resultado")
        eventos.append(evento)

    # Sort eventos by fecha (date) with most recent first
    from datetime import datetime

    def parse_date(date_str):
        try:
            return datetime.strptime(date_str, "%m/%d/%Y")
        except (ValueError, TypeError):
            # Return a very old date if parsing fails
            return datetime(1900, 1, 1)

    # First sort in ascending order (oldest first)
    eventos.sort(key=lambda x: parse_date(x.get("fecha")))

    # Then reverse the entire array to get descending order (newest first)
    eventos.reverse()

    # Replace the existing arrays with the new eventos array
    data["eventos"] = eventos
    del data["tramites"]
    del data["votaciones"]

    # Remove the documents array
    if "documents" in data:
        del data["documents"]

    # Log the structured data
    logger.info("Data structure collected from scraping:")
    logger.info(json.dumps(data, indent=2, ensure_ascii=False))

    # Save the JSON to a file for reference
    json_filepath = os.path.join(output_dir, "scraped_data.json")
    with open(json_filepath, 'w', encoding='utf-8') as json_file:
        json.dump(data, indent=2, ensure_ascii=False, fp=json_file)
    logger.info(f"Saved structured data to {json_filepath}")

    cleanup_debug_files()

    return data

def extract_text_from_docx(docx_path):
    """Extracts text from a .docx file."""
    try:
        doc = docx.Document(docx_path)
        full_text = []
        
        # Get text from document properties
        try:
            core_props = doc.core_properties
            if core_props.title:
                full_text.append(f"Title: {core_props.title}")
            if core_props.subject:
                full_text.append(f"Subject: {core_props.subject}")
            if core_props.author:
                full_text.append(f"Author: {core_props.author}")
            if core_props.keywords:
                full_text.append(f"Keywords: {core_props.keywords}")
            full_text.append("")  # Empty line after properties
        except:
            pass
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                full_text.append(" | ".join(row_text))
        
        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"Error reading docx {docx_path}: {e}")
        return ""

def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file."""
    try:
        text = extract_text(pdf_path)
        return text
    except Exception as e:
        logger.error(f"Error extracting text from {pdf_path}: {e}")
        return ""

def cleanup_debug_files():
    debug_files = ['page_source.html', 'error_screenshot.png']
    for i in range(5):
        debug_files.append(f'loading_attempt_{i+1}.png')
        
    for file in debug_files:
        if os.path.exists(file):
            try:
                os.remove(file)
                logger.info(f"Cleaned up debug file: {file}")
            except Exception as e:
                logger.warning(f"Could not clean up {file}: {e}")

# --- Main Execution ---
if __name__ == "__main__":
    print("PYTHON SCRIPT STARTED WITH ARGS:", sys.argv, file=sys.stderr)
    import sys
    import argparse
    
    # Setup argument parser for better command-line options
    parser = argparse.ArgumentParser(description="Scrape bill data from SUTRA website")
    parser.add_argument("url", help="URL of the bill page to scrape")
    parser.add_argument("--no-extract", action="store_true", help="Skip document downloading and text extraction")
    parser.add_argument("--output-dir", default="scraped_data", help="Directory to save scraped data")
    
    # Parse only known args to handle when called from Node.js
    args, unknown = parser.parse_known_args()
    
    url = args.url
    output_directory = args.output_dir
    skip_downloads = args.no_extract
    
    logger.info(f"Starting scraper for URL: {url}")
    logger.info(f"Skip document downloads: {skip_downloads}")
    
    # Modify the scrape_and_download function behavior based on the flag
    if skip_downloads:
        # Create a patched version of download_and_process_doc that just returns the URL info
        original_download_func = download_and_process_doc
        
        def patched_download_func(doc_info, output_dir, extract_text=False):
            # Just return the document info without downloading
            return {
                "link_url": doc_info["link_url"],
                "description": doc_info.get("description", "Document"),
                "downloaded": False,
                "text_extracted": False
            }
        
        # Apply the monkey patch
        import types
        download_and_process_doc = patched_download_func
    
    # Call the scraper
    result = scrape_and_download(url, output_directory)

    if "error" in result:
        logger.error(result["error"])
        # Still output the result as JSON even on error
        print(json.dumps(result))
    else:
        # Save the result to a JSON file
        with open("result.json", "w", encoding="utf-8") as f:
            json.dump(result, indent=2, ensure_ascii=False, fp=f)
        logger.info("Results saved to result.json")
        
        # IMPORTANT: Print the result as JSON for the Node.js server to parse
        print(json.dumps(result))
        
        # Also print summary to logger only, not to stdout
        logger.info("\nSCRAPER SUMMARY:")
        logger.info(f"Measure Number: {result.get('measure_number', 'Not found')}")
        logger.info(f"Title: {result.get('title', 'Not found')}")
        logger.info(f"Status: {result.get('status', 'Not found')}")
        logger.info(f"Events found: {len(result.get('eventos', []))}")
        logger.info(f"Commissions found: {len(result.get('comisiones', []))}")
        
        if skip_downloads:
            logger.info("NOTE: Document downloading was skipped due to --no-extract flag")